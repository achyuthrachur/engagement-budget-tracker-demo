from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from version import APP_VERSION  # noqa: E402

INDIGO = "011E41"
AMBER = "F5A800"
TEXT = "333333"
MID = "4F4F4F"
LIGHT = "E0E0E0"
WHITE = "FFFFFF"


def shade(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    fill = properties.find(qn("w:shd"))
    if fill is None:
        fill = OxmlElement("w:shd")
        properties.append(fill)
    fill.set(qn("w:fill"), color)


def cell_margins(cell, top=120, start=150, bottom=120, end=150):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_run(run, size=None, bold=None, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in (("Title", 28), ("Heading 1", 19), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INDIGO)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"Engagement Budget Tracker {APP_VERSION}  |  ")
    set_run(run, size=8, color=MID)
    add_page_number(footer)
    for footer_run in footer.runs[1:]:
        set_run(footer_run, size=8, color=MID)


def add_banner(document, label, subtitle=None):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, INDIGO)
    cell_margins(cell, top=220, start=260, bottom=220, end=260)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(label)
    set_run(run, size=12, bold=True, color=WHITE)
    if subtitle:
        run = paragraph.add_run(f"\n{subtitle}")
        set_run(run, size=9, color=WHITE)
    document.add_paragraph()


def add_callout(document, title, text, color=AMBER):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(0.12)
    table.columns[1].width = Inches(6.7)
    accent = table.cell(0, 0)
    body = table.cell(0, 1)
    shade(accent, color)
    shade(body, "F7F7F7")
    cell_margins(accent, 80, 20, 80, 20)
    cell_margins(body, 150, 190, 150, 190)
    paragraph = body.paragraphs[0]
    run = paragraph.add_run(title)
    set_run(run, size=10.5, bold=True, color=INDIGO)
    run = paragraph.add_run(f"\n{text}")
    set_run(run, size=10, color=TEXT)
    document.add_paragraph()


def add_steps(document, steps):
    for step in steps:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(step)


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_checklist(document, items):
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        run = paragraph.add_run(f"☐  {item}")
        set_run(run, size=10.5)


def add_data_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        shade(cell, INDIGO)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell)
        run = cell.paragraphs[0].add_run(label)
        set_run(run, size=9, bold=True, color=WHITE)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            shade(cells[index], "F7F7F7" if row_index % 2 else WHITE)
            cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            run = cells[index].paragraphs[0].add_run(str(value))
            set_run(run, size=9, color=TEXT)
    document.add_paragraph()
    return table


def add_section_title(document, number, title, description=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number:02d}")
    set_run(run, size=10, bold=True, color=AMBER)
    run = paragraph.add_run(f"  {title}")
    set_run(run, size=19, bold=True, color=INDIGO)
    if description:
        paragraph = document.add_paragraph(description)
        paragraph.paragraph_format.space_after = Pt(10)


def build_document(output_path):
    document = Document()
    configure_document(document)
    document.core_properties.title = "Engagement Budget Tracker instructions"
    document.core_properties.subject = "Installation, setup, weekly operation and recovery"
    document.core_properties.author = "Crowe"
    document.core_properties.keywords = "budget tracker, Cognos, engagement"

    add_banner(document, "Engagement Budget Tracker", "Installation and operating guide")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Simple guidance for a reliable engagement budget")
    set_run(run, size=28, bold=True, color=INDIGO)
    subtitle = document.add_paragraph(
        "Use this guide to open the tracker, set up an engagement, run the weekly "
        "budget process and recover safely from mistakes.")
    subtitle.paragraph_format.space_after = Pt(18)
    for run in subtitle.runs:
        set_run(run, size=13, color=MID)

    add_callout(
        document,
        "Recommended start",
        "Double-click B2A_Budget_Tracker.exe. The tracker selects an available "
        "local port and opens in your default browser. No Python installation or "
        "administrator access is required.")

    details = document.add_table(rows=3, cols=2)
    details.alignment = WD_TABLE_ALIGNMENT.LEFT
    details.style = "Table Grid"
    for row, (label, value) in enumerate((
            ("Version", APP_VERSION),
            ("Designed for", "Engagement teams and first-time technology users"),
            ("Data model", "Local Windows application with automatic recovery backups"))):
        details.cell(row, 0).text = label
        details.cell(row, 1).text = value
        shade(details.cell(row, 0), INDIGO)
        for run in details.cell(row, 0).paragraphs[0].runs:
            set_run(run, size=9, bold=True, color=WHITE)
        for run in details.cell(row, 1).paragraphs[0].runs:
            set_run(run, size=9, color=TEXT)
        cell_margins(details.cell(row, 0))
        cell_margins(details.cell(row, 1))

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(24)
    run = paragraph.add_run("Keep this document with the application when sharing it.")
    set_run(run, size=10, bold=True, color=INDIGO)
    document.add_page_break()

    add_section_title(document, 1, "Choose how to open the tracker")
    document.add_heading("Option A  Direct run", level=2)
    add_steps(document, [
        "Save B2A_Budget_Tracker.exe in a folder you can find.",
        "Double-click B2A_Budget_Tracker.exe.",
        "Wait for your default browser to open the Dashboard.",
        "If Windows displays an organizational security prompt, stop and contact support. "
        "Do not bypass company security controls.",
    ])
    add_callout(
        document,
        "Already running",
        "Double-clicking the executable again opens the existing tracker instead of "
        "creating a second database or duplicate session.")

    document.add_heading("Option B  Create a desktop shortcut", level=2)
    add_steps(document, [
        "Extract all files from the ZIP into one folder.",
        "Double-click install.bat.",
        "Select OK when the installation message appears.",
        "Use the Engagement Budget Tracker shortcut created on the desktop.",
    ])
    document.add_paragraph(
        "The installer does not require administrator access. Application updates do "
        "not overwrite the separate production database.")

    add_section_title(document, 2, "Complete first-time setup")
    add_steps(document, [
        "Open Settings from the left navigation.",
        "Review each role and its default internal rate.",
        "Enter engagement and contract discounts as normal percentages. For example, "
        "enter 10 for ten percent.",
        "Review the variance thresholds used to flag unusual weekly changes.",
        "Select Save settings.",
        "Select Create recovery backup.",
        "Return to the Dashboard and select I understand on the welcome card.",
    ])

    add_section_title(document, 3, "Create an engagement")
    add_steps(document, [
        "Select New engagement.",
        "Enter the exact Cognos project identifier as the engagement code.",
        "Enter the client name and engagement lead.",
        "Choose Simple for one overall budget or Complex for phase and weekly planning.",
        "For a Complex engagement, enter the first Monday and planned duration.",
        "Add every expected worker using the exact Cognos “Last, First” name.",
        "Review role, rates, planned hours and offshore designation for every worker.",
        "For a Complex engagement, add every phase and its signed statement of work fee.",
        "Set phase target hours and select Distribute target.",
        "Adjust the weekly cells until every reconciliation difference is zero.",
        "Review and select the baseline confirmation.",
        "Select Create engagement.",
    ])
    add_callout(
        document,
        "Baseline protection",
        "The engagement begins in Planning. The first committed Cognos import activates "
        "the engagement and locks baseline hours, rates and statement of work values. Later changes "
        "require a documented revision reason.")

    add_section_title(document, 4, "Run the weekly budget")
    add_checklist(document, [
        "Create a recovery backup in Settings.",
        "Export the raw Time and Cost Detail workbook from Cognos.",
        "Open the correct engagement and select Weekly import.",
        "Choose the Cognos file and select Preview import.",
        "Resolve unknown workers.",
        "Assign unmatched phases where appropriate.",
        "Leave project mismatches excluded unless independently verified.",
        "Review variance warnings.",
        "Review selected row, hour and contract-fee totals.",
        "Select Review and commit import.",
        "Open each phase and update future Forecast values.",
        "Review Overview and Export the partner report.",
    ])
    add_callout(
        document,
        "Safe import sequence",
        "Preview never changes the budget. Commit creates a recovery backup before "
        "writing time entries. Duplicate transactions remain excluded.")

    document.add_heading("Import warnings", level=2)
    add_data_table(document, ["Warning", "What it means", "Action"], [
        ("Duplicate", "The transaction was already imported", "Leave excluded"),
        ("Zero hours", "The row contains no time", "Leave excluded"),
        ("Unknown worker", "No active team member matches the worker", "Add or reactivate the worker"),
        ("Project mismatch", "The row project identifier differs from the engagement", "Verify independently before selecting"),
        ("Unmatched phase", "The Cognos description is not assigned", "Assign a tracker phase"),
        ("Variance review", "The weekly change exceeds a threshold", "Confirm the change is reasonable"),
    ], widths=[1.2, 2.5, 2.6])

    add_section_title(document, 5, "Review results and update forecasts")
    add_bullets(document, [
        "Overview shows engagement status, budget status, statement of work budget, actual fees and realization.",
        "Phase detail shows Budget, Actual and Forecast by person and week.",
        "Budget is the approved baseline and cannot be edited directly after activation.",
        "Actual comes from committed Cognos imports.",
        "Forecast is the future estimate and should be reviewed after each import.",
        "Crowe-paid expenses reduce realization. Client-paid expenses are informational.",
        "An approved budget addition increases the engagement budget. An approved budget reduction decreases it. Change orders add to a phase.",
    ])

    add_section_title(document, 6, "Correct a mistake")
    document.add_heading("Bad weekly import", level=2)
    add_steps(document, [
        "Open History for the engagement.",
        "Find the affected snapshot by week-ending date.",
        "Select Delete and confirm the warning.",
        "Preview and commit the corrected Cognos file.",
    ])
    document.add_paragraph(
        "The tracker creates a recovery backup before deleting a snapshot.")

    document.add_heading("Larger recovery", level=2)
    add_steps(document, [
        "Open Settings.",
        "Under Database and recovery, choose a known-good .db backup.",
        "Select Validate and restore.",
        "Review the confirmation and continue.",
        "Return to the Dashboard and verify the expected engagements.",
    ])
    add_callout(
        document,
        "Restore validation",
        "The tracker checks SQLite integrity and required application tables before "
        "replacing the current database. It preserves the current database first.")

    add_section_title(document, 7, "Close or reopen an engagement")
    add_steps(document, [
        "Open the engagement Overview.",
        "Use the lifecycle control to choose Close engagement or Reopen engagement.",
        "Enter a clear reason, including who approved the decision.",
        "Confirm the change.",
    ])
    document.add_paragraph(
        "Closed engagements are read-only. Reopening restores controlled editing and "
        "retains the lifecycle event in History.")

    add_section_title(document, 8, "Understand the terms")
    add_data_table(document, ["Term", "Plain-language meaning"], [
        ("Statement of work budget", "The signed fee budget for the work"),
        ("Standard rate", "The internal value of a person’s time"),
        ("Engagement rate", "The rate used for planned engagement fees"),
        ("Contract rate", "The Cognos rate compared with the statement of work budget"),
        ("Advance billing tracking", "Informational tracking that does not enforce the budget"),
        ("Realization", "Actual contract fees to date less Crowe-paid expenses, divided by actual standard fees to date"),
        ("Approved budget addition", "An approved engagement-wide budget increase"),
        ("Approved budget reduction", "An approved budget decrease"),
        ("Change order", "An approved addition assigned to a phase"),
        ("Budget", "The approved baseline"),
        ("Actual", "Committed Cognos time"),
        ("Forecast", "The future estimate"),
    ], widths=[1.7, 4.8])

    add_section_title(document, 9, "Find data and backups")
    add_data_table(document, ["Item", "Windows location"], [
        ("Installed application", r"%LOCALAPPDATA%\Crowe\B2A Budget Tracker\App"),
        ("Production database", r"%LOCALAPPDATA%\Crowe\B2A Budget Tracker\App\budget_tracker.db"),
        ("Automatic backups", r"%LOCALAPPDATA%\Crowe\B2A Budget Tracker\App\Backups"),
    ], widths=[2.1, 4.4])
    document.add_paragraph(
        "The tracker retains the 20 most recent automatic backups. Database files may "
        "contain confidential engagement information and must only be transferred "
        "through approved channels.")

    add_section_title(document, 10, "Troubleshoot common issues")
    add_data_table(document, ["Issue", "What to do"], [
        ("Browser does not open", "Double-click the executable once more. If needed, restart Windows and retry."),
        ("Windows security prompt", "Stop and contact support. Do not bypass organizational controls."),
        ("Ports are in use", "Close other local applications or restart Windows."),
        ("Worker is unknown", "Add or reactivate the exact Cognos worker name in Team and budget."),
        ("Project mismatch", "Confirm that the file belongs to the engagement before selecting the row."),
        ("Wrong phase totals", "Assign unmatched descriptions and verify the phase mapping."),
        ("Application update", "Save work, close the browser and run install.bat from the newer ZIP."),
    ], widths=[2.0, 4.5])

    document.add_page_break()
    add_banner(document, "Administrator and support notes")
    add_section_title(document, 11, "Prepare a release for sharing")
    add_steps(document, [
        "Run build.bat from the project folder.",
        f"Code-sign release\\B2A_Budget_Tracker.exe using the approved organizational certificate.",
        "Recreate the versioned ZIP from the five release files after signing so it "
        "contains the signed executable. Do not rerun build.bat because a rebuild "
        "replaces the signed file.",
        "Distribute the ZIP or executable only through an approved internal channel.",
        "Ask a first-time user to complete one observed setup and weekly-import pilot.",
    ])
    add_callout(
        document,
        "Code signing",
        "The development build is functional but should be signed before broad "
        "distribution. Signing reduces security warnings and proves publisher identity.")

    document.add_heading("Update behavior", level=2)
    add_bullets(document, [
        "Direct EXE use stores data outside the executable.",
        "The optional installer stops an existing tracker before replacing the application.",
        "The database and Backups folder remain untouched during application upgrades.",
        "The executable binds only to 127.0.0.1 and is not exposed to the network.",
    ])

    document.add_heading("Support information to collect", level=2)
    add_bullets(document, [
        "Windows version",
        "Application version shown in Settings",
        "Schema version shown in Settings",
        "Exact error message",
        "Whether the issue occurred during preview or commit",
        "A database backup transferred only through an approved confidential-data channel",
    ])

    add_callout(
        document,
        "Confidentiality",
        "Never request Cognos time data or a tracker database through email or an "
        "unapproved file-transfer method.",
        color=INDIGO)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    destination = (Path(sys.argv[1]) if len(sys.argv) > 1
                   else ROOT / "release" / "B2A_Budget_Tracker_Instructions.docx")
    result = build_document(destination)
    print(f"Built {result}")
